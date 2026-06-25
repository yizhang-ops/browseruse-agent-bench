#!/usr/bin/env python3
"""Build embedded DOCX and self-contained HTML report for generation comparison."""

from __future__ import annotations

import base64
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = REPO_ROOT / "reports"
ASSET_DIR = REPORT_DIR / "assets"
FIGURE = ASSET_DIR / "generation_failure_comparison.png"
DOCX_OUT = REPORT_DIR / "doubao_glm_generation_comparison.docx"
HTML_OUT = REPORT_DIR / "doubao_glm_generation_comparison.html"

BLUE = "355F9F"
BLUE_LIGHT = "E8EEF7"
AMBER = "C6922E"
TEAL = "3E8582"
INK = "1C2B33"
MUTED = "5A6472"
BORDER = "D7DBE2"
LATIN_FONT = "Arial Unicode MS"
EAST_ASIA_FONT = "Arial Unicode MS"

OVERALL = [
    ["Doubao Seed 2.0 Pro", "130 / 210", "80", "61.90%", "80"],
    ["Doubao Seed 2.1 Pro", "99 / 210", "111", "47.14%", "111"],
    ["GLM-5.1", "143 / 210", "67", "68.10%", "67"],
    ["GLM-5.2", "139 / 210", "71", "66.19%", "70"],
]

DOUBAO_LAYER = [
    ["Task Reasoning", "26", "32.5%", "21", "18.9%", "-5"],
    ["Action Execution", "18", "22.5%", "57", "51.4%", "+39"],
    ["Web Constraints", "36", "45.0%", "33", "29.7%", "-3"],
]
DOUBAO_TYPES = [
    ["M1.1 Requirement Following", "18", "21", "+3"],
    ["M1.2 Target Selection", "4", "0", "-4"],
    ["M1.3 Evidence Grounding", "4", "0", "-4"],
    ["M2.1 UI Misoperation", "6", "4", "-2"],
    ["M2.2 Infinite Loop", "10", "40", "+30"],
    ["M2.3 Format Breakdown", "2", "13", "+11"],
    ["M3.1 Bot Defense", "15", "13", "-2"],
    ["M3.2 Access Barrier", "3", "2", "-1"],
    ["M3.3 Site Limitation", "18", "18", "0"],
]

GLM_LAYER = [
    ["Task Reasoning", "15", "22.4%", "27", "38.6%", "+12"],
    ["Action Execution", "17", "25.4%", "20", "28.6%", "+3"],
    ["Web Constraints", "35", "52.2%", "23", "32.9%", "-12"],
]
GLM_TYPES = [
    ["M1.1 Requirement Following", "10", "15", "+5"],
    ["M1.2 Target Selection", "3", "8", "+5"],
    ["M1.3 Evidence Grounding", "2", "4", "+2"],
    ["M2.1 UI Misoperation", "5", "4", "-1"],
    ["M2.2 Infinite Loop", "10", "12", "+2"],
    ["M2.3 Format Breakdown", "2", "4", "+2"],
    ["M3.1 Bot Defense", "15", "11", "-4"],
    ["M3.2 Access Barrier", "1", "0", "-1"],
    ["M3.3 Site Limitation", "19", "12", "-7"],
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 80, start: int = 90, bottom: int = 80, end: int = 90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def format_paragraph(paragraph, size: float = 10.5, bold: bool = False, color: str = INK) -> None:
    for run in paragraph.runs:
        run.font.name = LATIN_FONT
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor.from_string(color)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)


def add_para(doc: Document, text: str, size: float = 10.0, after: int = 5, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    run.font.name = LATIN_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(INK)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.font.name = LATIN_FONT
    run.font.bold = True
    run.font.size = Pt(15 if level == 1 else 12.5)
    run.font.color.rgb = RGBColor.from_string(BLUE if level == 1 else INK)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.1)
    run = p.add_run(text)
    run.font.name = LATIN_FONT
    run.font.size = Pt(9.8)
    run.font.color.rgb = RGBColor.from_string(INK)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = Inches(widths[i])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, BLUE_LIGHT)
        set_cell_border(cell)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i else WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(header)
        format_paragraph(p, size=9.2, bold=True, color=INK)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            cell.width = Inches(widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(value)
            color = INK
            if value.startswith("+"):
                color = AMBER
            elif value.startswith("-"):
                color = TEAL
            format_paragraph(p, size=8.8, bold=i == 0, color=color)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal.font.size = Pt(10.0)
    normal.font.color.rgb = RGBColor.from_string(INK)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("Doubao 与 GLM 在 LexBench-Browser 上的代际对比")
    r.font.name = LATIN_FONT
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)

    add_para(
        doc,
        "这份报告对比了两组模型代际更新在相同 BrowserUse scaffold 下的表现，评测集合为 LexBench-Browser All split。",
        size=10.4,
    )
    add_bullet(doc, "Doubao Seed 2.0 Pro vs. Doubao Seed 2.1 Pro")
    add_bullet(doc, "GLM-5.1 vs. GLM-5.2")
    add_para(
        doc,
        "任务完成率由 LexJudge 使用 gpt-4.1 评估；failure attribution 使用 LexBench failure taxonomy，并由 gpt-5.5 作为 judge。GLM-5.2 中有 1 条失败来自对危险请求的正确安全拒答，这类样本不属于 agent capability failure，因此已从错误归因统计中排除。",
        size=10.0,
    )

    fig = doc.add_paragraph()
    fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig.paragraph_format.space_before = Pt(4)
    fig.paragraph_format.space_after = Pt(3)
    fig.add_run().add_picture(str(FIGURE), width=Inches(7.0))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    rr = cap.add_run("Figure 1. Generation-level failure attribution comparison.")
    rr.font.name = LATIN_FONT
    rr.font.size = Pt(9)
    rr.font.italic = True
    rr.font.color.rgb = RGBColor.from_string(MUTED)

    add_heading(doc, "Overall Performance")
    add_table(
        doc,
        ["Model", "Successful Tasks", "Raw Failed", "Success Rate", "Attributed Failures"],
        OVERALL,
        [2.25, 1.35, 1.05, 1.1, 1.35],
    )
    add_para(
        doc,
        "Doubao Seed 2.1 Pro 相比 2.0 Pro 出现明显退化：少完成 31 个任务，Success Rate 下降 14.76 个百分点。退化并不主要来自更困难的网站环境，而是集中在 action-level failures，尤其是 Infinite Loop 和 Format Breakdown。",
    )
    add_para(
        doc,
        "GLM-5.2 与 GLM-5.1 的差距较小，但 raw task completion 仍略低：少完成 4 个任务，Success Rate 低 1.91 个百分点。如果把安全拒答样本从失败侧去掉，GLM-5.2 的非错误结果为 140/210，即 66.67%。",
    )

    doc.add_page_break()
    add_heading(doc, "Doubao 2.0 Pro vs. Doubao 2.1 Pro")
    add_table(
        doc,
        ["Error Layer", "2.0 Count", "2.0 Ratio", "2.1 Count", "2.1 Ratio", "Change"],
        DOUBAO_LAYER,
        [1.75, 1.05, 1.05, 1.05, 1.05, 0.85],
    )
    add_table(
        doc,
        ["Error Type", "Doubao 2.0", "Doubao 2.1", "Change"],
        DOUBAO_TYPES,
        [3.85, 1.05, 1.05, 0.9],
    )
    add_para(
        doc,
        "Doubao Seed 2.1 Pro 的核心问题是 operational stability。M2.2 Infinite Loop 从 10 条上升到 40 条，M2.3 Format Breakdown 从 2 条上升到 13 条。相比之下，Web Constraints 的绝对数量没有增加，M3.3 Site Limitation 仍然是 18 条。",
    )
    add_para(
        doc,
        "这说明 2.1 的退化主要是 model-side，而不是 environment-side。它不是更频繁地被网站阻断，而是更容易无法维持有效的交互循环、无法从页面状态中恢复，或者无法保持可解析的 action format。",
    )

    doc.add_page_break()
    add_heading(doc, "GLM-5.1 vs. GLM-5.2")
    add_table(
        doc,
        ["Error Layer", "5.1 Count", "5.1 Ratio", "5.2 Count", "5.2 Ratio", "Change"],
        GLM_LAYER,
        [1.75, 1.05, 1.05, 1.05, 1.05, 0.85],
    )
    add_table(
        doc,
        ["Error Type", "GLM-5.1", "GLM-5.2", "Change"],
        GLM_TYPES,
        [3.85, 1.05, 1.05, 0.9],
    )
    add_para(
        doc,
        "GLM-5.2 在 website-side constraints 上更干净：Web Constraints 从 35 条下降到 23 条，其中 Bot Defense 和 Site Limitation 都减少了。",
    )
    add_para(
        doc,
        "但是，Web Constraints 的减少没有转化为更高的整体任务完成率。错误更多地转移到 model-side categories：Task Reasoning 从 15 条上升到 27 条，主要集中在 Requirement Following 和 Target Selection。",
    )

    add_heading(doc, "Takeaways")
    add_bullet(doc, "Doubao 2.1 Pro 变差的主要原因是 Action Execution 不稳定。")
    add_bullet(doc, "GLM-5.2 虽然更少受到 Web Constraints 影响，但 model-side reasoning errors 增加，因此整体略低于 GLM-5.1。")
    add_bullet(doc, "这套 taxonomy 能把 model capability failures 和 website-side constraints 分开，从而更清楚地定位代际变化的来源。")

    doc.save(DOCX_OUT)


def html_table(headers: list[str], rows: list[list[str]]) -> str:
    head = "".join(f"<th>{h}</th>" for h in headers)
    body_rows = []
    for row in rows:
        body_rows.append("<tr>" + "".join(f"<td>{v}</td>" for v in row) + "</tr>")
    return f"<table><thead><tr>{head}</tr></thead><tbody>{''.join(body_rows)}</tbody></table>"


def build_html() -> None:
    image_b64 = base64.b64encode(FIGURE.read_bytes()).decode("ascii")
    html = f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Doubao 与 GLM 在 LexBench-Browser 上的代际对比</title>
  <style>
    :root {{
      --blue: #{BLUE};
      --amber: #{AMBER};
      --teal: #{TEAL};
      --ink: #{INK};
      --muted: #{MUTED};
      --border: #{BORDER};
      --fill: #{BLUE_LIGHT};
    }}
    body {{
      margin: 0;
      background: #f6f8fb;
      color: var(--ink);
      font-family: -apple-system, BlinkMacSystemFont, "PingFang SC", "Segoe UI", Arial, sans-serif;
      line-height: 1.55;
    }}
    main {{
      max-width: 1120px;
      margin: 0 auto;
      padding: 48px 36px 64px;
      background: white;
      box-shadow: 0 16px 48px rgba(28, 43, 51, 0.08);
      min-height: 100vh;
    }}
    h1 {{
      margin: 0 0 18px;
      color: var(--blue);
      font-size: 34px;
      line-height: 1.18;
    }}
    h2 {{
      margin: 36px 0 12px;
      color: var(--blue);
      font-size: 24px;
    }}
    p {{ margin: 0 0 14px; }}
    ul {{ margin: 0 0 18px 22px; padding: 0; }}
    li {{ margin: 4px 0; }}
    figure {{ margin: 28px 0 30px; }}
    figure img {{
      width: 100%;
      height: auto;
      display: block;
      border: 1px solid var(--border);
      background: white;
    }}
    figcaption {{
      margin-top: 8px;
      color: var(--muted);
      font-style: italic;
      font-size: 14px;
    }}
    table {{
      width: 100%;
      border-collapse: collapse;
      margin: 12px 0 22px;
      font-size: 14px;
    }}
    th, td {{
      border: 1px solid var(--border);
      padding: 8px 10px;
      vertical-align: middle;
    }}
    th {{
      background: var(--fill);
      text-align: left;
      font-weight: 700;
    }}
    td:not(:first-child), th:not(:first-child) {{ text-align: center; }}
    .lead {{
      font-size: 17px;
      color: #263846;
      margin-bottom: 18px;
    }}
    .note {{
      border-left: 4px solid var(--blue);
      background: #f4f7fb;
      padding: 12px 14px;
      margin: 18px 0 22px;
      color: #2d3a45;
    }}
    .taxonomy {{
      display: grid;
      grid-template-columns: repeat(3, 1fr);
      gap: 12px;
      margin: 16px 0 26px;
    }}
    .tax-card {{
      border: 1px solid var(--border);
      border-top: 4px solid var(--blue);
      border-radius: 8px;
      padding: 12px 14px;
      background: #ffffff;
    }}
    .tax-card:nth-child(2) {{ border-top-color: var(--amber); }}
    .tax-card:nth-child(3) {{ border-top-color: var(--teal); }}
    .tax-card strong {{
      display: block;
      margin-bottom: 6px;
      font-size: 15px;
    }}
    .tax-card span {{
      color: var(--muted);
      font-size: 13px;
    }}
    .tax-card dl {{
      margin: 10px 0 0;
      display: grid;
      gap: 5px;
      font-size: 12.5px;
    }}
    .tax-card div {{
      display: grid;
      grid-template-columns: 54px 1fr;
      gap: 7px;
    }}
    .tax-card dt {{
      font-weight: 700;
      color: var(--ink);
    }}
    .tax-card dd {{
      margin: 0;
      color: #3f4b57;
    }}
    .takeaways {{
      border-left: 4px solid var(--amber);
      background: #fff8e9;
      padding: 12px 16px;
    }}
    code {{
      background: #eef2f7;
      padding: 1px 4px;
      border-radius: 4px;
    }}
    @media (max-width: 800px) {{
      main {{ padding: 32px 18px 48px; }}
      .taxonomy {{ grid-template-columns: 1fr; }}
      table {{ font-size: 12px; }}
      th, td {{ padding: 6px 7px; }}
    }}
  </style>
</head>
<body>
<main>
  <h1>Doubao 与 GLM 在 LexBench-Browser 上的代际对比</h1>
  <p class="lead">这份报告对比了两组模型代际更新在相同 BrowserUse scaffold 下的表现，评测集合为 LexBench-Browser <code>All</code> split。</p>
  <ul>
    <li>Doubao Seed 2.0 Pro vs. Doubao Seed 2.1 Pro</li>
    <li>GLM-5.1 vs. GLM-5.2</li>
  </ul>
  <p class="note">任务完成率由 LexJudge 使用 <code>gpt-4.1</code> 评估；failure attribution 使用 LexBench failure taxonomy，并由 <code>gpt-5.5</code> 作为 judge。GLM-5.2 中有 1 条失败来自对危险请求的正确安全拒答，这类样本不属于 agent capability failure，因此已从错误归因统计中排除。</p>

  <h2>Failure Taxonomy</h2>
  <div class="taxonomy">
    <div class="tax-card">
      <strong>M1 · Task Reasoning</strong>
      <span>模型没有正确理解、保持或验证任务要求，例如 Requirement Following、Target Selection、Evidence Grounding。</span>
      <dl>
        <div><dt>M1.1</dt><dd>Requirement Following：遗漏或偏离任务约束。</dd></div>
        <div><dt>M1.2</dt><dd>Target Selection：选错页面、对象、商品或结果。</dd></div>
        <div><dt>M1.3</dt><dd>Evidence Grounding：答案缺少足够页面证据支撑。</dd></div>
      </dl>
    </div>
    <div class="tax-card">
      <strong>M2 · Action Execution</strong>
      <span>模型在浏览器交互或工具调用层面失稳，例如 UI Misoperation、Infinite Loop、Format Breakdown。</span>
      <dl>
        <div><dt>M2.1</dt><dd>UI Misoperation：点击、输入、导航等浏览器操作错误。</dd></div>
        <div><dt>M2.2</dt><dd>Infinite Loop：重复无效动作，无法推进任务。</dd></div>
        <div><dt>M2.3</dt><dd>Format Breakdown：输出或 action format 无法被框架解析。</dd></div>
      </dl>
    </div>
    <div class="tax-card">
      <strong>M3 · Web Constraints</strong>
      <span>失败主要来自网站或浏览环境限制，例如 Bot Defense、Access Barrier、Site Limitation。</span>
      <dl>
        <div><dt>M3.1</dt><dd>Bot Defense：验证码、风控或反自动化机制阻断。</dd></div>
        <div><dt>M3.2</dt><dd>Access Barrier：登录、权限、地区或账号状态限制。</dd></div>
        <div><dt>M3.3</dt><dd>Site Limitation：网站功能、内容或页面状态本身不可用。</dd></div>
      </dl>
    </div>
  </div>

  <figure>
    <img src="data:image/png;base64,{image_b64}" alt="Generation-level failure attribution comparison">
    <figcaption>Figure 1. Generation-level failure attribution comparison.</figcaption>
  </figure>

  <h2>Overall Performance</h2>
  {html_table(["Model", "Successful Tasks", "Raw Failed Tasks", "Success Rate", "Attributed Failures"], OVERALL)}
  <p>Doubao Seed 2.1 Pro 相比 2.0 Pro 出现明显退化：少完成 31 个任务，Success Rate 下降 14.76 个百分点。退化并不主要来自更困难的网站环境，而是集中在 action-level failures，尤其是 Infinite Loop 和 Format Breakdown。</p>
  <p>GLM-5.2 与 GLM-5.1 的差距较小，但 raw task completion 仍略低：少完成 4 个任务，Success Rate 低 1.91 个百分点。如果把安全拒答样本从失败侧去掉，GLM-5.2 的非错误结果为 140/210，即 66.67%。</p>

  <h2>Doubao 2.0 Pro vs. Doubao 2.1 Pro</h2>
  {html_table(["Error Layer", "Doubao 2.0 Count", "Doubao 2.0 Ratio", "Doubao 2.1 Count", "Doubao 2.1 Ratio", "Change"], DOUBAO_LAYER)}
  {html_table(["Error Type", "Doubao 2.0", "Doubao 2.1", "Change"], DOUBAO_TYPES)}
  <p>Doubao Seed 2.1 Pro 的核心问题是 operational stability。<code>M2.2 Infinite Loop</code> 从 10 条上升到 40 条，<code>M2.3 Format Breakdown</code> 从 2 条上升到 13 条。相比之下，Web Constraints 的绝对数量没有增加，<code>M3.3 Site Limitation</code> 仍然是 18 条。</p>
  <p>这说明 2.1 的退化主要是 model-side，而不是 environment-side。它不是更频繁地被网站阻断，而是更容易无法维持有效的交互循环、无法从页面状态中恢复，或者无法保持可解析的 action format。</p>

  <h2>GLM-5.1 vs. GLM-5.2</h2>
  {html_table(["Error Layer", "GLM-5.1 Count", "GLM-5.1 Ratio", "GLM-5.2 Count", "GLM-5.2 Ratio", "Change"], GLM_LAYER)}
  {html_table(["Error Type", "GLM-5.1", "GLM-5.2", "Change"], GLM_TYPES)}
  <p>GLM-5.2 在 website-side constraints 上更干净：Web Constraints 从 35 条下降到 23 条，其中 Bot Defense 和 Site Limitation 都减少了。</p>
  <p>但是，Web Constraints 的减少没有转化为更高的整体任务完成率。错误更多地转移到 model-side categories：Task Reasoning 从 15 条上升到 27 条，主要集中在 Requirement Following 和 Target Selection。</p>

  <h2>Takeaways</h2>
  <div class="takeaways">
    <ul>
      <li>Doubao 2.1 Pro 变差的主要原因是 Action Execution 不稳定。</li>
      <li>GLM-5.2 虽然更少受到 Web Constraints 影响，但 model-side reasoning errors 增加，因此整体略低于 GLM-5.1。</li>
      <li>这套 taxonomy 能把 model capability failures 和 website-side constraints 分开，从而更清楚地定位代际变化的来源。</li>
    </ul>
  </div>
</main>
</body>
</html>
"""
    HTML_OUT.write_text(html, encoding="utf-8")


def main() -> None:
    build_docx()
    build_html()
    print(DOCX_OUT)
    print(HTML_OUT)


if __name__ == "__main__":
    main()
